"""
Blog routes for Opinian platform
Handles blog post creation, editing, and viewing
"""

import os
import logging
import re
from io import BytesIO
from flask import Blueprint, render_template, request, redirect, url_for, flash, session, jsonify, send_file
from werkzeug.utils import secure_filename
from datetime import datetime
import psycopg2
from psycopg2.extras import RealDictCursor
from app import get_db_connection, login_required, role_required, allowed_file, log_user_activity
from ai_service import ai_service
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

bp = Blueprint('blog', __name__, url_prefix='/blog')

def get_upload_path():
    """Get upload path for blog images"""
    return os.path.join('uploads', 'blog_images')

@bp.route('/')
def blog_index():
    """Blog index page - list all published blog posts"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            
            # Get all published blog posts from active groups
            cursor.execute("""
                SELECT bp.*, u.username, u.first_name, u.last_name, u.profile_image_url, g.name as group_name
                FROM blog_posts bp
                JOIN users u ON bp.author_id = u.id
                LEFT JOIN groups g ON bp.group_id = g.id
                WHERE bp.is_published = TRUE AND (g.is_active = TRUE OR bp.group_id IS NULL)
                ORDER BY bp.published_at DESC
            """)
            blog_posts = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            return render_template('blog/index.html', blog_posts=blog_posts)
        else:
            flash('Database connection error', 'danger')
            return render_template('blog/index.html', blog_posts=[])
            
    except Exception as e:
        flash('Error loading blog posts', 'danger')
        return render_template('blog/index.html', blog_posts=[])

@bp.route('/post/<slug>')
def view_post(slug):
    """View a single blog post"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)

            # Get blog post
            cursor.execute("""
                SELECT bp.*, u.username, u.first_name, u.last_name, u.profile_image_url, u.bio, g.name as group_name
                FROM blog_posts bp
                JOIN users u ON bp.author_id = u.id
                LEFT JOIN groups g ON bp.group_id = g.id
                WHERE bp.slug = %s AND bp.is_published = TRUE
            """, (slug,))

            post = cursor.fetchone()

            if not post:
                flash('Blog post not found', 'danger')
                return redirect(url_for('blog.blog_index'))

            # Increment view count
            cursor.execute("UPDATE blog_posts SET view_count = view_count + 1 WHERE id = %s", (post['id'],))
            conn.commit()

            # Get related posts (same group or same tags)
            cursor.execute("""
                SELECT id, title, slug, published_at, excerpt
                FROM blog_posts
                WHERE group_id = %s AND id != %s AND is_published = TRUE
                ORDER BY published_at DESC
                LIMIT 5
            """, (post['group_id'], post['id']))
            related_posts = cursor.fetchall()

            # Get comments for this post (approved only, with nested replies)
            cursor.execute("""
                SELECT c.*, u.username, u.first_name, u.last_name, u.profile_image_url
                FROM comments c
                JOIN users u ON c.user_id = u.id
                WHERE c.blog_post_id = %s AND c.is_approved = TRUE AND c.is_deleted = FALSE
                ORDER BY c.created_at ASC
            """, (post['id'],))
            all_comments = cursor.fetchall()

            # Organize comments into tree structure
            comments = []
            comment_map = {}
            for comment in all_comments:
                comment['replies'] = []
                comment_map[comment['id']] = comment
                if comment['parent_id'] is None:
                    comments.append(comment)
                else:
                    parent = comment_map.get(comment['parent_id'])
                    if parent:
                        parent['replies'].append(comment)

            # Get comment count
            comment_count = len(all_comments)

            cursor.close()
            conn.close()

            return render_template('blog/view.html', post=post, related_posts=related_posts,
                                   comments=comments, comment_count=comment_count)
        else:
            flash('Database connection error', 'danger')
            return redirect(url_for('blog.blog_index'))

    except Exception as e:
        flash('Error loading blog post', 'danger')
        return redirect(url_for('blog.blog_index'))

@bp.route('/create', methods=['GET', 'POST'])
@login_required
def create_post():
    """Create a new blog post"""
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        excerpt = request.form.get('excerpt')
        tags = request.form.get('tags', '')
        meta_description = request.form.get('meta_description')
        meta_keywords = request.form.get('meta_keywords')

        # Check action button (publish or draft)
        action = request.form.get('action', 'draft')

        # Moderation Logic based on user role (per REQUIREMENTS.md)
        user_role = session.get('user_role')
        needs_moderation = False
        is_published = False

        if action == 'publish':
            # SuperAdmin, Admin, SuperUser can publish directly
            if user_role in ['SuperAdmin', 'Admin', 'SuperUser']:
                is_published = True
            else:
                # Regular User role - needs moderation
                needs_moderation = True
                is_published = False  # Will be set to True upon approval

        # Handle publish scheduling
        publish_type = request.form.get('publish_type', 'immediate')
        scheduled_date = request.form.get('publish_date')

        # Determine published_at timestamp
        published_at = None
        if is_published:
            if publish_type == 'scheduled' and scheduled_date:
                try:
                    # Parse the scheduled date
                    from dateutil import parser
                    published_at = parser.parse(scheduled_date)
                except:
                    published_at = datetime.utcnow()
            else:
                published_at = datetime.utcnow()
        
        # Handle file upload
        featured_image_url = None
        if 'featured_image' in request.files:
            file = request.files['featured_image']
            if file and file.filename and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                upload_path = get_upload_path()
                os.makedirs(upload_path, exist_ok=True)
                
                # Generate unique filename
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                unique_filename = f"{timestamp}_{filename}"
                file_path = os.path.join(upload_path, unique_filename)
                
                file.save(file_path)
                featured_image_url = file_path
        
        try:
            conn = get_db_connection()
            if conn:
                cursor = conn.cursor()
                
                # Generate slug from title
                import re
                slug = re.sub(r'[^a-zA-Z0-9-]+', '-', title.lower()).strip('-')
                
                # Ensure unique slug
                cursor.execute("SELECT id FROM blog_posts WHERE slug = %s", (slug,))
                if cursor.fetchone():
                    slug = f"{slug}-{int(datetime.now().timestamp())}"
                
                # Insert blog post
                cursor.execute("""
                    INSERT INTO blog_posts
                    (title, slug, content, excerpt, author_id, group_id, featured_image_url,
                     tags, meta_description, meta_keywords, is_published, published_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    RETURNING id
                """, (
                    title, slug, content, excerpt, session['user_id'], session.get('group_id'),
                    featured_image_url, tags.split(',') if tags else [], meta_description, meta_keywords,
                    is_published, published_at
                ))
                
                post_id = cursor.fetchone()[0]

                # Add to moderation queue if needed
                if needs_moderation:
                    cursor.execute("""
                        INSERT INTO moderation_queue (content_type, content_id, status, created_at)
                        VALUES (%s, %s, %s, %s)
                    """, ('blog_post', post_id, 'pending', datetime.utcnow()))
                    logger.info(f"Blog post {post_id} added to moderation queue")

                conn.commit()
                cursor.close()
                conn.close()

                # Log activity
                log_user_activity(session['user_id'], 'create_blog_post', 'blog_post', post_id)

                # Show appropriate message
                if needs_moderation:
                    flash('Blog post submitted for moderation. It will be published after admin approval.', 'info')
                elif is_published:
                    flash('Blog post published successfully!', 'success')
                else:
                    flash('Blog post saved as draft!', 'success')

                if is_published and not needs_moderation:
                    return redirect(url_for('blog.view_post', slug=slug))
                else:
                    return redirect(url_for('blog.my_posts'))
            else:
                flash('Database connection error', 'danger')
                
        except Exception as e:
            flash('Error creating blog post', 'danger')
            logger.error(f"Error creating blog post: {e}")
    
    return render_template('blog/create.html')

@bp.route('/edit/<int:post_id>', methods=['GET', 'POST'])
@login_required
def edit_post(post_id):
    """Edit an existing blog post"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            
            # Get blog post
            cursor.execute("""
                SELECT * FROM blog_posts WHERE id = %s
            """, (post_id,))
            post = cursor.fetchone()
            
            if not post:
                flash('Blog post not found', 'danger')
                return redirect(url_for('blog.my_posts'))
            
            # Check permissions
            if session['user_role'] not in ['SuperAdmin', 'Admin'] and post['author_id'] != session['user_id']:
                flash('You do not have permission to edit this post', 'danger')
                return redirect(url_for('blog.my_posts'))
            
            if request.method == 'POST':
                title = request.form.get('title')
                content = request.form.get('content')
                excerpt = request.form.get('excerpt')
                tags = request.form.get('tags', '')
                meta_description = request.form.get('meta_description')
                meta_keywords = request.form.get('meta_keywords')

                # Check action button (publish or draft)
                action = request.form.get('action', 'draft')
                is_published = (action == 'publish')

                # Handle publish scheduling
                publish_type = request.form.get('publish_type', 'immediate')
                scheduled_date = request.form.get('publish_date')

                # Determine published_at timestamp
                published_at = post['published_at']  # Keep existing if already set
                if is_published:
                    if publish_type == 'scheduled' and scheduled_date:
                        try:
                            from dateutil import parser
                            published_at = parser.parse(scheduled_date)
                        except:
                            published_at = datetime.utcnow()
                    elif not published_at:  # Only set if not already published
                        published_at = datetime.utcnow()
                else:
                    published_at = None  # Clear if saving as draft
                
                # Handle file upload
                featured_image_url = post['featured_image_url']
                if 'featured_image' in request.files:
                    file = request.files['featured_image']
                    if file and file.filename and allowed_file(file.filename):
                        # Delete old image if exists
                        if featured_image_url and os.path.exists(featured_image_url):
                            os.remove(featured_image_url)
                        
                        filename = secure_filename(file.filename)
                        upload_path = get_upload_path()
                        os.makedirs(upload_path, exist_ok=True)
                        
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        unique_filename = f"{timestamp}_{filename}"
                        file_path = os.path.join(upload_path, unique_filename)
                        
                        file.save(file_path)
                        featured_image_url = file_path
                
                # Update slug if title changed
                import re
                slug = re.sub(r'[^a-zA-Z0-9-]+', '-', title.lower()).strip('-')
                if slug != post['slug']:
                    cursor.execute("SELECT id FROM blog_posts WHERE slug = %s AND id != %s", (slug, post_id))
                    if cursor.fetchone():
                        slug = f"{slug}-{int(datetime.now().timestamp())}"
                
                # Update blog post
                cursor.execute("""
                    UPDATE blog_posts
                    SET title = %s, slug = %s, content = %s, excerpt = %s,
                        featured_image_url = %s, tags = %s, meta_description = %s,
                        meta_keywords = %s, is_published = %s, published_at = %s,
                        updated_at = %s
                    WHERE id = %s
                """, (
                    title, slug, content, excerpt, featured_image_url,
                    tags.split(',') if tags else [],
                    meta_description, meta_keywords, is_published, published_at,
                    datetime.utcnow(), post_id
                ))
                
                conn.commit()
                cursor.close()
                conn.close()
                
                # Log activity
                log_user_activity(session['user_id'], 'edit_blog_post', 'blog_post', post_id)
                
                flash('Blog post updated successfully!', 'success')
                if is_published:
                    return redirect(url_for('blog.view_post', slug=slug))
                else:
                    return redirect(url_for('blog.my_posts'))
            
            cursor.close()
            conn.close()
            
            return render_template('blog/edit.html', post=post)
        else:
            flash('Database connection error', 'danger')
            return redirect(url_for('blog.my_posts'))
            
    except Exception as e:
        flash('Error loading blog post', 'danger')
        logger.error(f"Error editing blog post: {e}")
        return redirect(url_for('blog.my_posts'))

@bp.route('/delete/<int:post_id>', methods=['POST'])
@login_required
def delete_post(post_id):
    """Delete a blog post"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            
            # Get blog post
            cursor.execute("SELECT * FROM blog_posts WHERE id = %s", (post_id,))
            post = cursor.fetchone()
            
            if not post:
                flash('Blog post not found', 'danger')
                return redirect(url_for('blog.my_posts'))
            
            # Check permissions
            if session['user_role'] not in ['SuperAdmin', 'Admin'] and post['author_id'] != session['user_id']:
                flash('You do not have permission to delete this post', 'danger')
                return redirect(url_for('blog.my_posts'))
            
            # Delete featured image if exists
            if post['featured_image_url'] and os.path.exists(post['featured_image_url']):
                os.remove(post['featured_image_url'])
            
            # Delete blog post
            cursor.execute("DELETE FROM blog_posts WHERE id = %s", (post_id,))
            conn.commit()
            cursor.close()
            conn.close()
            
            # Log activity
            log_user_activity(session['user_id'], 'delete_blog_post', 'blog_post', post_id)
            
            flash('Blog post deleted successfully!', 'success')
            return redirect(url_for('blog.my_posts'))
        else:
            flash('Database connection error', 'danger')
            return redirect(url_for('blog.my_posts'))
            
    except Exception as e:
        flash('Error deleting blog post', 'danger')
        logger.error(f"Error deleting blog post: {e}")
        return redirect(url_for('blog.my_posts'))

@bp.route('/my-posts')
@login_required
def my_posts():
    """List user's blog posts"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            
            user_role = session['user_role']
            user_id = session['user_id']
            
            if user_role in ['SuperAdmin', 'Admin']:
                # Admins can see all posts in their group
                group_id = session.get('group_id')
                if group_id is not None:
                    cursor.execute("""
                        SELECT bp.*, u.username
                        FROM blog_posts bp
                        JOIN users u ON bp.author_id = u.id
                        WHERE bp.group_id = %s
                        ORDER BY bp.created_at DESC
                    """, (group_id,))
                else:
                    cursor.execute("""
                        SELECT bp.*, u.username
                        FROM blog_posts bp
                        JOIN users u ON bp.author_id = u.id
                        WHERE bp.group_id IS NULL
                        ORDER BY bp.created_at DESC
                    """)
            else:
                # Regular users can only see their own posts
                cursor.execute("""
                    SELECT * FROM blog_posts 
                    WHERE author_id = %s
                    ORDER BY created_at DESC
                """, (user_id,))
            
            blog_posts = cursor.fetchall()
            cursor.close()
            conn.close()
            
            return render_template('blog/my_posts.html', blog_posts=blog_posts)
        else:
            flash('Database connection error', 'danger')
            return render_template('blog/my_posts.html', blog_posts=[])
            
    except Exception as e:
        flash('Error loading blog posts', 'danger')
        logger.error(f"Error loading my posts: {e}")
        return render_template('blog/my_posts.html', blog_posts=[])

@bp.route('/ai-assistant')
@login_required
def ai_assistant():
    """AI writing assistant page"""
    return render_template('blog/ai_assistant.html')

@bp.route('/ai-generate', methods=['POST'])
@login_required
def ai_generate():
    """Generate or improve content using AI for the embedded modal"""
    try:
        data = request.get_json()
        input_text = data.get('input')
        style = data.get('style', 'professional')
        action = data.get('action', 'generate')

        if not input_text:
            return jsonify({
                'success': False,
                'error': 'Input text is required'
            }), 400

        # Build prompt based on action and style
        if action == 'improve':
            prompt = f"Improve the following content in a {style} style:\n\n{input_text}"
        else:
            prompt = f"Write a blog post in a {style} style based on this idea:\n\n{input_text}"

        # Use AI service to generate content
        result = ai_service.generate_blog_content(prompt, 'blog_post')

        # Log activity
        log_user_activity(
            session['user_id'],
            f'ai_{action}_content',
            'blog_content',
            None,
            {'style': style, 'action': action}
        )

        if result.get('success'):
            return jsonify(result)
        else:
            return jsonify(result), 500

    except Exception as e:
        logger.error(f"Error in AI generate: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to generate content',
            'message': str(e)
        }), 500

@bp.route('/generate-content', methods=['POST'])
@login_required
def generate_content():
    """Generate content using AI"""
    try:
        data = request.get_json()
        prompt = data.get('prompt')
        content_type = data.get('content_type', 'blog_post')

        if not prompt:
            return jsonify({
                'success': False,
                'error': 'Prompt is required'
            }), 400

        # Use AI service to generate content
        result = ai_service.generate_blog_content(prompt, content_type)

        # Log activity
        log_user_activity(
            session['user_id'],
            'generate_ai_content',
            'blog_content',
            None,
            {'prompt': prompt[:100], 'content_type': content_type}
        )

        if result.get('success'):
            return jsonify(result)
        else:
            return jsonify(result), 500

    except Exception as e:
        logger.error(f"Error generating content: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to generate content',
            'message': str(e)
        }), 500

@bp.route('/generate-titles', methods=['POST'])
@login_required
def generate_titles():
    """Generate title suggestions using AI"""
    try:
        data = request.get_json()
        topic = data.get('topic')
        count = data.get('count', 5)

        if not topic:
            return jsonify({
                'success': False,
                'error': 'Topic is required'
            }), 400

        # Use AI service to generate titles
        result = ai_service.generate_title_suggestions(topic, count)

        # Log activity
        log_user_activity(
            session['user_id'],
            'generate_ai_titles',
            'blog_titles',
            None,
            {'topic': topic[:100]}
        )

        if result.get('success'):
            return jsonify(result)
        else:
            return jsonify(result), 500

    except Exception as e:
        logger.error(f"Error generating titles: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to generate titles',
            'message': str(e)
        }), 500

@bp.route('/improve-content', methods=['POST'])
@login_required
def improve_content():
    """Improve existing content using AI"""
    try:
        data = request.get_json()
        content = data.get('content')
        instructions = data.get('instructions', 'Improve the writing quality and engagement')

        if not content:
            return jsonify({
                'success': False,
                'error': 'Content is required'
            }), 400

        # Use AI service to improve content
        result = ai_service.improve_content(content, instructions)

        # Log activity
        log_user_activity(
            session['user_id'],
            'improve_ai_content',
            'blog_content',
            None,
            {'instructions': instructions[:100]}
        )

        if result.get('success'):
            return jsonify(result)
        else:
            return jsonify(result), 500

    except Exception as e:
        logger.error(f"Error improving content: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to improve content',
            'message': str(e)
        }), 500

@bp.route('/generate-excerpt', methods=['POST'])
@login_required
def generate_excerpt():
    """Generate excerpt from full content using AI"""
    try:
        data = request.get_json()
        content = data.get('content')
        max_length = data.get('max_length', 200)

        if not content:
            return jsonify({
                'success': False,
                'error': 'Content is required'
            }), 400

        # Use AI service to generate excerpt
        result = ai_service.generate_excerpt(content, max_length)

        # Log activity
        log_user_activity(
            session['user_id'],
            'generate_ai_excerpt',
            'blog_excerpt',
            None
        )

        if result.get('success'):
            return jsonify(result)
        else:
            return jsonify(result), 500

    except Exception as e:
        logger.error(f"Error generating excerpt: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to generate excerpt',
            'message': str(e)
        }), 500


# ============== COMMENT ROUTES ==============

@bp.route('/post/<int:post_id>/comment', methods=['POST'])
@login_required
def add_comment(post_id):
    """Add a comment to a blog post"""
    try:
        content = request.form.get('content', '').strip()
        parent_id = request.form.get('parent_id')

        if not content:
            flash('Comment cannot be empty', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

        if len(content) > 5000:
            flash('Comment is too long (max 5000 characters)', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)

            # Verify post exists and is published
            cursor.execute("""
                SELECT id, slug, author_id FROM blog_posts
                WHERE id = %s AND is_published = TRUE
            """, (post_id,))
            post = cursor.fetchone()

            if not post:
                flash('Blog post not found', 'danger')
                cursor.close()
                conn.close()
                return redirect(url_for('blog.blog_index'))

            # Validate parent_id if provided
            if parent_id:
                cursor.execute("""
                    SELECT id FROM comments
                    WHERE id = %s AND blog_post_id = %s AND is_deleted = FALSE
                """, (parent_id, post_id))
                if not cursor.fetchone():
                    parent_id = None

            # Insert comment
            cursor.execute("""
                INSERT INTO comments (blog_post_id, user_id, parent_id, content)
                VALUES (%s, %s, %s, %s)
                RETURNING id
            """, (post_id, session['user_id'], parent_id if parent_id else None, content))

            comment_id = cursor.fetchone()['id']
            conn.commit()

            # Log activity
            log_user_activity(session['user_id'], 'add_comment', 'comment', comment_id,
                              {'post_id': post_id, 'is_reply': bool(parent_id)})

            cursor.close()
            conn.close()

            flash('Comment added successfully!', 'success')
            return redirect(url_for('blog.view_post', slug=post['slug']) + f'#comment-{comment_id}')
        else:
            flash('Database connection error', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

    except Exception as e:
        logger.error(f"Error adding comment: {e}")
        flash('Error adding comment', 'danger')
        return redirect(request.referrer or url_for('blog.blog_index'))


@bp.route('/comment/<int:comment_id>/edit', methods=['POST'])
@login_required
def edit_comment(comment_id):
    """Edit a comment"""
    try:
        content = request.form.get('content', '').strip()

        if not content:
            flash('Comment cannot be empty', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)

            # Get comment and verify ownership
            cursor.execute("""
                SELECT c.*, bp.slug FROM comments c
                JOIN blog_posts bp ON c.blog_post_id = bp.id
                WHERE c.id = %s AND c.is_deleted = FALSE
            """, (comment_id,))
            comment = cursor.fetchone()

            if not comment:
                flash('Comment not found', 'danger')
                cursor.close()
                conn.close()
                return redirect(url_for('blog.blog_index'))

            # Check permission (owner or admin)
            if comment['user_id'] != session['user_id'] and session['user_role'] not in ['SuperAdmin', 'Admin']:
                flash('You do not have permission to edit this comment', 'danger')
                cursor.close()
                conn.close()
                return redirect(url_for('blog.view_post', slug=comment['slug']))

            # Update comment
            cursor.execute("""
                UPDATE comments SET content = %s, updated_at = %s
                WHERE id = %s
            """, (content, datetime.utcnow(), comment_id))
            conn.commit()

            # Log activity
            log_user_activity(session['user_id'], 'edit_comment', 'comment', comment_id)

            cursor.close()
            conn.close()

            flash('Comment updated successfully!', 'success')
            return redirect(url_for('blog.view_post', slug=comment['slug']) + f'#comment-{comment_id}')
        else:
            flash('Database connection error', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

    except Exception as e:
        logger.error(f"Error editing comment: {e}")
        flash('Error editing comment', 'danger')
        return redirect(request.referrer or url_for('blog.blog_index'))


@bp.route('/comment/<int:comment_id>/delete', methods=['POST'])
@login_required
def delete_comment(comment_id):
    """Delete a comment (soft delete)"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor(cursor_factory=RealDictCursor)

            # Get comment details
            cursor.execute("""
                SELECT c.*, bp.slug, bp.author_id as post_author_id FROM comments c
                JOIN blog_posts bp ON c.blog_post_id = bp.id
                WHERE c.id = %s AND c.is_deleted = FALSE
            """, (comment_id,))
            comment = cursor.fetchone()

            if not comment:
                flash('Comment not found', 'danger')
                cursor.close()
                conn.close()
                return redirect(url_for('blog.blog_index'))

            # Check permission (comment owner, post author, or admin)
            can_delete = (
                comment['user_id'] == session['user_id'] or
                comment['post_author_id'] == session['user_id'] or
                session['user_role'] in ['SuperAdmin', 'Admin']
            )

            if not can_delete:
                flash('You do not have permission to delete this comment', 'danger')
                cursor.close()
                conn.close()
                return redirect(url_for('blog.view_post', slug=comment['slug']))

            # Soft delete
            cursor.execute("""
                UPDATE comments SET is_deleted = TRUE, updated_at = %s
                WHERE id = %s
            """, (datetime.utcnow(), comment_id))
            conn.commit()

            # Log activity
            log_user_activity(session['user_id'], 'delete_comment', 'comment', comment_id)

            cursor.close()
            conn.close()

            flash('Comment deleted successfully!', 'success')
            return redirect(url_for('blog.view_post', slug=comment['slug']))
        else:
            flash('Database connection error', 'danger')
            return redirect(request.referrer or url_for('blog.blog_index'))

    except Exception as e:
        logger.error(f"Error deleting comment: {e}")
        flash('Error deleting comment', 'danger')
        return redirect(request.referrer or url_for('blog.blog_index'))


# ============== EXPORT TO WORD ROUTES ==============

def strip_html_tags(text):
    """Remove HTML tags from text"""
    if not text:
        return ""
    # Remove HTML tags
    clean = re.compile('<.*?>')
    text = re.sub(clean, '', text)
    # Replace common HTML entities
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&amp;', '&')
    text = text.replace('&lt;', '<')
    text = text.replace('&gt;', '>')
    text = text.replace('&quot;', '"')
    text = text.replace('&#39;', "'")
    return text.strip()


def html_to_docx_paragraphs(doc, html_content):
    """Convert HTML content to Word document paragraphs with basic formatting"""
    if not html_content:
        return

    # Split by common HTML tags
    # Handle headings
    h1_pattern = re.compile(r'<h1[^>]*>(.*?)</h1>', re.DOTALL | re.IGNORECASE)
    h2_pattern = re.compile(r'<h2[^>]*>(.*?)</h2>', re.DOTALL | re.IGNORECASE)
    h3_pattern = re.compile(r'<h3[^>]*>(.*?)</h3>', re.DOTALL | re.IGNORECASE)
    p_pattern = re.compile(r'<p[^>]*>(.*?)</p>', re.DOTALL | re.IGNORECASE)
    li_pattern = re.compile(r'<li[^>]*>(.*?)</li>', re.DOTALL | re.IGNORECASE)
    br_pattern = re.compile(r'<br\s*/?>', re.IGNORECASE)

    # Replace br tags with newlines
    html_content = br_pattern.sub('\n', html_content)

    # Process headings
    for match in h1_pattern.finditer(html_content):
        text = strip_html_tags(match.group(1))
        if text:
            heading = doc.add_heading(text, level=1)

    for match in h2_pattern.finditer(html_content):
        text = strip_html_tags(match.group(1))
        if text:
            heading = doc.add_heading(text, level=2)

    for match in h3_pattern.finditer(html_content):
        text = strip_html_tags(match.group(1))
        if text:
            heading = doc.add_heading(text, level=3)

    # Process paragraphs
    for match in p_pattern.finditer(html_content):
        text = strip_html_tags(match.group(1))
        if text:
            doc.add_paragraph(text)

    # Process list items
    for match in li_pattern.finditer(html_content):
        text = strip_html_tags(match.group(1))
        if text:
            doc.add_paragraph(text, style='List Bullet')

    # If no structured content found, add as plain text
    if not (h1_pattern.search(html_content) or h2_pattern.search(html_content) or
            h3_pattern.search(html_content) or p_pattern.search(html_content)):
        clean_text = strip_html_tags(html_content)
        if clean_text:
            for para in clean_text.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())


@bp.route('/export-post/<int:post_id>')
@login_required
def export_post_to_word(post_id):
    """Export a blog post to Word document"""
    try:
        conn = get_db_connection()
        if not conn:
            flash('Database connection error', 'danger')
            return redirect(url_for('blog.my_posts'))

        cursor = conn.cursor(cursor_factory=RealDictCursor)

        # Get blog post
        cursor.execute("""
            SELECT bp.*, u.username, u.first_name, u.last_name, g.name as group_name
            FROM blog_posts bp
            JOIN users u ON bp.author_id = u.id
            LEFT JOIN groups g ON bp.group_id = g.id
            WHERE bp.id = %s
        """, (post_id,))
        post = cursor.fetchone()

        cursor.close()
        conn.close()

        if not post:
            flash('Blog post not found', 'danger')
            return redirect(url_for('blog.my_posts'))

        # Check permission
        if session['user_role'] not in ['SuperAdmin', 'Admin'] and post['author_id'] != session['user_id']:
            flash('You do not have permission to export this post', 'danger')
            return redirect(url_for('blog.my_posts'))

        # Create Word document
        doc = Document()

        # Add title
        title = doc.add_heading(post['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(
            f"By {post['first_name']} {post['last_name']}\n"
            f"{post['published_at'].strftime('%B %d, %Y') if post['published_at'] else 'Draft'}"
        )
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        if post['group_name']:
            org_para = doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_run = org_para.add_run(post['group_name'])
            org_run.font.size = Pt(10)
            org_run.font.italic = True

        doc.add_paragraph()  # Spacer

        # Add featured image if exists
        if post['featured_image_url'] and os.path.exists(post['featured_image_url']):
            try:
                doc.add_picture(post['featured_image_url'], width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacer
            except Exception as e:
                logger.error(f"Error adding image to Word doc: {e}")

        # Add excerpt if exists
        if post['excerpt']:
            excerpt_para = doc.add_paragraph()
            excerpt_run = excerpt_para.add_run(strip_html_tags(post['excerpt']))
            excerpt_run.font.italic = True
            excerpt_run.font.size = Pt(11)
            doc.add_paragraph()  # Spacer

        # Add content
        html_to_docx_paragraphs(doc, post['content'])

        # Add tags if exist
        if post['tags']:
            doc.add_paragraph()
            tags_para = doc.add_paragraph()
            tags_run = tags_para.add_run(f"Tags: {', '.join(post['tags'])}")
            tags_run.font.size = Pt(9)
            tags_run.font.color.rgb = RGBColor(100, 100, 100)

        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated with Opinian")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)

        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Generate filename
        safe_title = re.sub(r'[^\w\s-]', '', post['title']).strip().replace(' ', '_')
        filename = f"{safe_title}_{post['id']}.docx"

        # Log activity
        log_user_activity(session['user_id'], 'export_post_to_word', 'blog_post', post_id)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"Error exporting blog post to Word: {e}")
        flash('Error exporting blog post', 'danger')
        return redirect(url_for('blog.my_posts'))


@bp.route('/export-content', methods=['POST'])
@login_required
def export_content_to_word():
    """Export AI-generated content to Word document"""
    try:
        data = request.get_json()
        title = data.get('title', 'Untitled Document')
        content = data.get('content', '')

        if not content:
            return jsonify({
                'success': False,
                'error': 'No content to export'
            }), 400

        # Create Word document
        doc = Document()

        # Add title
        title_heading = doc.add_heading(title, level=0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_para.add_run(f"Generated on {datetime.utcnow().strftime('%B %d, %Y')}")
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Add content
        html_to_docx_paragraphs(doc, content)

        # Add footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated with Opinian AI Assistant")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(150, 150, 150)

        # Save to BytesIO
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # Generate filename
        safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
        filename = f"{safe_title}.docx"

        # Log activity
        log_user_activity(session['user_id'], 'export_ai_content_to_word', 'ai_content', None)

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        logger.error(f"Error exporting AI content to Word: {e}")
        return jsonify({
            'success': False,
            'error': 'Failed to export content',
            'message': str(e)
        }), 500